"""Generate the sprite redraw inventory as a .docx with creative suggestions.

Outputs to ./SPRITE_INVENTORY.docx in the repo root. Run with:
    python scripts/generate-sprite-list-docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def add_table(doc, rows):
    """rows[0] is header. Each subsequent row is [emoji, name, where, suggestion]."""
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    widths = [Inches(0.55), Inches(1.6), Inches(2.0), Inches(2.7)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(11) if i > 0 else Pt(10)
            if i == 0:
                run.bold = True
                shade_cell(cell, '2E2E5C')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif j == 0:
                run.font.size = Pt(18)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold


def main():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # ── Title ──────────────────────────────────────────
    title = doc.add_heading('GregSweeper sprite redraw inventory', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    body(doc, 'Generated 2026-05-07. Each sprite below is currently rendered as an '
              'emoji in the codebase. The "Idea direction" column is a starting '
              'point — you can draw anything that fits the slot. Recommended export '
              'format: PNG with transparent alpha, 128x128 for cell content / power-ups '
              '/ medals, 256x256 for mode cards, 64-96px for nav buttons. SVG also '
              'fine if you prefer vector. File names go into /assets/sprites/{slug}.png '
              'and the asset-loader (TBD) handles the rest.')
    doc.add_paragraph()

    # ── Tier 1 ─────────────────────────────────────────
    heading(doc, 'Tier 1 — Cell sprites (default theme)')
    body(doc, 'On screen during every second of gameplay. Replacing these alone '
              'removes ~80% of the "feels like an emoji-stitched MVP" impression. '
              'Each at 128x128, transparent alpha. The smiley triplet should feel '
              'like one character (we currently use the default emoji set; consider '
              'making "Greg" your mascot — a marine biologist with safety goggles, '
              'lab coat, or pickaxe.)', bold=False)
    add_table(doc, [
        ['Icon', 'Asset', 'Where it appears', 'Idea direction'],
        ['💣', 'Mine', 'Loss screen, daily bomb hit',
         'Antique pin-cluster mine with sparking fuse; or a sleeker mechanical '
         'detonator. Could also be a marine mine (round with horns) tying into your '
         'biology background. Should look DANGEROUS at small cell sizes.'],
        ['🚩', 'Flag', 'Player-flagged cells',
         'Triangular pennant on a wooden pole (classic minesweeper). Or warning '
         'stake with reflective stripes. Or a small surveying flag — keep silhouette '
         'unambiguous against the cell background.'],
        ['😊', 'Smiley (idle)', 'Reset button (top center)',
         'Greg as the mascot — friendly, neutral expression. Could be helmet + '
         'goggles + slight smile. Single sprite reused across themes.'],
        ['😎', 'Smiley (win)', 'Reset button after win',
         'Same Greg, sunglasses + grin. Or fist pump. Should read as "victory" '
         'instantly even at small button size.'],
        ['😵', 'Smiley (loss)', 'Reset button after loss / daily bomb',
         'Greg with X-eyes / dazed expression / soot-covered. Tonally "oh no" '
         'rather than "horror movie."'],
        ['💥', 'Strike', 'Daily/weekly bomb-hit popup, leaderboard column header',
         'Comic-book BOOM with debris. Or a more graphic explosion plume. Smaller '
         'than the mine — appears inline with text in leaderboard headers, so needs '
         'to read at ~16px.'],
    ])

    # ── Tier 2 ─────────────────────────────────────────
    heading(doc, 'Tier 2 — Modifier overlays')
    body(doc, 'In-cell visual identifiers for modifier mechanics. Strong '
              'silhouettes more important than detail — must read at 30-40px cell '
              'sizes on phone screens. Each modifier should be visually distinct '
              "from the others (don't make Sonar and Compass look similar).")
    add_table(doc, [
        ['Icon', 'Asset', 'Modifier mechanic', 'Idea direction'],
        ['🧱', 'Wall', 'Walls modifier (currently shown as edge lines, but icon '
         'used in legend / help)',
         'Brick texture or stone-block pattern. Used in help/legend UI, not '
         'in-cell. Could mirror the actual edge-line aesthetic we use on the board.'],
        ['🤥', 'Liar', 'Liar cells (displayed number is off by ±1)',
         'A cell number with an arrow / flicker / scratch through it indicating '
         '"don\'t trust me." Or a face with crossed fingers. Or a question mark '
         'over a digit. Avoid Pinocchio nose (the emoji default) — too literal.'],
        ['❓', 'Mystery', 'Hidden-number cells (number not displayed)',
         'A blurred / fogged-out number. Or a magnifying glass over a question '
         'mark. Or a swirling fog cloud. Should feel "I know something is here '
         'but the count is hidden."'],
        ['🔒', 'Locked', 'Locked cells (unrevealable until safe neighbors revealed)',
         'A simple padlock — but distinguish from the "Lv. X locked" achievement '
         'lock. Could be a chain link, a bolted hatch, or a vault door. Should '
         'visually invite "tap me later."'],
        ['🌀', 'Wormhole', 'Paired cells where revealing one teleports the other',
         'Spiraling vortex / portal / two linked rings. Each pair gets the SAME '
         'sprite color or numbered tag so player can match. Consider drawing '
         'multiple pair-color variants.'],
        ['🪞', 'Mirror', 'Paired cells that swap displayed counts',
         'Mirror frame with a glint of reflection. Or two arrows curving into '
         'each other. Avoid "mirror as in glass" — needs to read as "swap."'],
        ['🔴', 'Pressure plate', 'Plate cells with countdown timer',
         'A floor plate / button / circular metal disc. Should look like '
         'something you stand on, not just "red dot." A small countdown digit '
         'overlays it in code, so leave space for that.'],
        ['📡', 'Sonar', 'Cells that report mine count over a 5×5 area',
         'A radar dish / sonar wave fan / pulsing rings. Currently rendered as '
         'emoji + number prefix; if you draw it, make sure a number can sit '
         'next to or inside it cleanly.'],
        ['🧭', 'Compass', 'Cells with directional arrow pointing toward nearest mine',
         'Compass rose with prominent needle direction. Currently rendered as '
         'emoji + arrow direction (N/S/E/W/NE/etc.) — your sprite should support '
         'rotating the arrow indicator (or pre-draw 8 rotations).'],
        ['💨', 'Mine shift', 'Chaos-mode indicator that mines just moved',
         'Motion lines / wind streaks / arrows showing displacement. Or a small '
         'whoosh effect. Used briefly as a flash, not a persistent cell state.'],
    ])

    # ── Tier 3 ─────────────────────────────────────────
    heading(doc, 'Tier 3 — Mode cards')
    body(doc, 'Giant icons on the title screen. Recommended 256x256. Should feel '
              'cohesive as a SET — matched line weight, consistent illustrative '
              'voice. Different palettes per card to help players associate visual '
              'with mode at a glance.')
    add_table(doc, [
        ['Icon', 'Mode', 'What it is', 'Idea direction'],
        ['⛏️', 'Challenge', '120-level main mode',
         'A pickaxe striking a stone tile. Or layered mineshaft cross-section '
         'showing depth/levels. Or Greg with hardhat. Should feel "the main '
         'thing you do."'],
        ['⏱️', 'Quick Play', '4-tier timed mode (Beginner → Extreme)',
         'A stopwatch with motion lines. Or a runner silhouette + clock. '
         'Tonally "race" — should feel snappy, not laborious.'],
        ['📅', 'Daily', "Today's puzzle, one shot",
         'A calendar page with today\'s date highlighted. Or a sun cresting '
         'over a single cell. Avoid the obvious "calendar" emoji literal — '
         'something that says "every day a new one." Could even be a coffee '
         'cup + cell, "your morning ritual."'],
        ['🎁', 'Bonus Daily', 'One-off bonus dailies on special dates',
         'A wrapped present with a bow. But CURRENTLY also used for the '
         'Collection footer button — these need to differ. Bonus daily could '
         'use a confetti burst or "extra" badge instead.'],
        ['🏁', 'Weekly', "This week's puzzle, 7 attempts on the same board",
         'A checkered flag (currently). Or a 7-day calendar grid showing all '
         'attempts. Or a lap-counter with stripes. Should feel "longer-form '
         'commitment" than daily.'],
        ['🌀', 'Chaos', 'Endless modifier-stacking rapid mode',
         'A chaotic vortex / glitch effect / scrambling tiles. Same emoji as '
         'wormhole — these MUST differ. Chaos is more aggressive: motion blur, '
         'cracked tiles, lightning. Wormhole is more "physics."'],
        ['🎓', 'Skill Trainer', 'Tutorial mode (currently hidden, code intact)',
         'Graduation cap is fine. Or a hand pointing at a board. Or a magnifying '
         'glass over a 1-2-3 tutorial cell. Reserve for re-launch.'],
    ])

    # ── Tier 4 ─────────────────────────────────────────
    heading(doc, 'Tier 4 — Power-ups')
    body(doc, 'In the power-up bar at the bottom of every game. 128x128. The bar '
              'is small — silhouettes must be readable at ~32px rendered size. '
              'Color-coding helps (e.g., shield = blue, lifeline = red).')
    add_table(doc, [
        ['Icon', 'Asset', 'What it does', 'Idea direction'],
        ['🔍', 'Reveal Safe', 'Reveals one safe cell',
         'A magnifying glass with a checkmark inside. Or a hand with a glowing '
         'finger. Could also be a torch/flashlight beam. Should feel "safe '
         'guidance."'],
        ['🛡️', 'Shield', 'Blocks one mine hit',
         'A traditional kite shield, or sci-fi energy shield. Differentiate '
         'from achievements that also use shield iconography — could add a '
         'spark/hit-mark for the in-game power-up version.'],
        ['🎯', 'Scan', 'Shows mine count for a row or column',
         'A target reticle with crosshairs / scanning line. Or a row/column '
         'highlight effect. Should feel "I am inspecting this line."'],
        ['❤️', 'Lifeline', 'Auto-saves you from one mine',
         'A heart with a pulse line, or a phone-a-friend cord, or a parachute. '
         '"Lifeline" is the right metaphor — something that catches you. Avoid '
         'the plain heart, it gets confused with health.'],
        ['🧲', 'Magnet', 'Pulls mines from a 3×3 area to board edges',
         'Classic horseshoe magnet with motion lines / pulled mine sprites. '
         'Could include the 3×3 hint as part of the design.'],
        ['🔬', 'X-Ray', 'Reveals 5×5 area\'s mines',
         'X-ray view of a hand/grid showing mines as dark spots. Or a microscope '
         '(currently emoji default) but X-ray is more apt — change the metaphor. '
         "Currently same emoji as the Diagnostics button — these should differ."],
    ])

    # ── Tier 5 ─────────────────────────────────────────
    heading(doc, 'Tier 5 — Speed rating medals')
    body(doc, 'End-of-game and leaderboard. 128x128, tightly cohesive set. Standard '
              'medal hierarchy works (bronze → silver → gold → diamond / platinum). '
              'These also double as achievement tier markers, so worth getting right.')
    add_table(doc, [
        ['Icon', 'Asset', 'Tier', 'Idea direction'],
        ['💎', 'Diamond', 'Top tier (fastest)',
         'A multi-faceted gem catching light, or a polished obsidian shard. '
         'Could even be a cluster of small gems for the "above and beyond" feel. '
         'Differentiate from typical "gem" emoji — make it FEEL prestigious.'],
        ['🥇', 'Gold', '2nd tier',
         'Classic gold medal on ribbon. The set should share a frame/ribbon '
         'style and only differ in metal/color.'],
        ['🥈', 'Silver', '3rd tier',
         'Same frame, silver. Slightly less ornate / fewer rays / smaller star.'],
        ['🥉', 'Bronze', '4th tier (slowest qualifying)',
         'Same frame, bronze. Should still feel earned, not consolation.'],
    ])

    # ── Tier 6 ─────────────────────────────────────────
    heading(doc, 'Tier 6 — Navigation / footer / nav-bar buttons')
    body(doc, 'Smaller sprites — these go in 36-48px buttons. 64x64 or 96x96 PNG, '
              'or SVG. Minimum-detail style. The current emoji set has multiple '
              'duplicates that need to differentiate (Collection vs Bonus Daily, '
              'Help vs Mystery, Diagnostics vs X-Ray, What\'s New vs Share). '
              'Treat these as a pictogram set with a unified line weight.')
    add_table(doc, [
        ['Icon', 'Asset', 'Where', 'Idea direction'],
        ['🏠', 'Home', 'In-game nav bar',
         'Simple house silhouette. Standard.'],
        ['📊', 'Stats', 'Title + nav',
         'Bar chart or trend line. Should feel data-y, not "dashboard busy."'],
        ['🏅', 'Achievements', 'Title + nav',
         'A medal pinned to ribbon, simpler than the speed-rating set. Could '
         'use a unique color (gold-ish but distinct from the gold medal sprite).'],
        ['🏆', 'Leaderboard', 'Title + nav',
         'Trophy cup with a small "1" or laurel. Distinct from achievement medal.'],
        ['🎁', 'Collection',
         'Title + nav (CURRENTLY DUPLICATE of Bonus Daily emoji)',
         'A stack of cards / inventory grid / treasure chest. Anything BUT '
         'a wrapped gift. The collection screen is where unlocked themes / '
         'effects live — could be a paint palette or a row of icons.'],
        ['❓', 'Help', 'Nav (CURRENTLY DUPLICATE of Mystery cell emoji)',
         'A speech bubble with "?" inside. Or an "i" info icon. NOT the bare '
         'question mark — must differ from the mystery-cell sprite.'],
        ['🔊', 'Unmuted', 'Nav toggle (currently)',
         'Speaker with sound waves. Pair with the muted version below.'],
        ['🔇', 'Muted', 'Nav toggle (currently)',
         'Speaker with X / strike-through. Same shape as unmuted, modified.'],
        ['⚙️', 'Settings', 'Title + nav',
         'Standard gear. Don\'t overthink.'],
        ['📋', 'What\'s New', "Title footer (CURRENTLY DUPLICATE of Share emoji)",
         'A scroll, or a megaphone, or a "release notes" page-and-star. '
         'Differentiate from Share.'],
        ['📋', 'Share', 'Game-over screen',
         'Send/forward arrow, or three-dot share node, or a paper airplane. '
         'Distinct from What\'s New.'],
        ['🔄', 'Check for Updates', 'Settings → Check for Updates',
         'Circular arrow / refresh icon. Standard.'],
        ['🗑️', 'Reset Profile', 'Settings → Reset Profile',
         'Trash can with a lid being opened. Or a recycle arrow into trash. '
         'Should feel "destructive but recoverable" — not too alarming.'],
        ['🔬', 'Diagnostics', "Settings → Diagnostics (CURRENTLY DUPLICATE of X-Ray "
         'power-up)',
         'A clipboard with checkmarks, or a stethoscope, or a debug-bug icon. '
         'Differentiate from X-Ray power-up.'],
        ['🔔', 'Notifications on', 'Settings + toasts',
         'Bell with sound waves. Pair with the off version.'],
        ['🔕', 'Notifications off', 'Settings + toasts',
         'Bell with strike-through. Same shape as on, modified.'],
        ['▶️', 'Resume', 'Title screen "Resume Game" button',
         'Play triangle. Simple.'],
        ['🔥', 'Streak fire', 'Title daily-card label, headers',
         'A small flame. Used inline with text — must read at 16-20px. Could '
         'be more abstract (like a flame-shaped chevron) since it appears next '
         'to a number ("🔥 5 day streak").'],
        ['👋', 'Wave (welcome)', 'Onboarding modal header',
         'A waving hand, or a confetti burst, or "Hi!" speech bubble. '
         'Tonally welcoming.'],
        ['👆', 'Tap-pointer', 'Flag-mode indicator',
         'A pointing finger / tap gesture. Used to indicate "tap to flag" '
         'mode. Could be a finger + small flag combo.'],
        ['⚠️', 'Warning', 'Temporary-mode toast',
         'Yellow triangle with !. Used when localStorage isn\'t available '
         '(private browsing). Standard warning glyph is fine.'],
        ['🟢', 'Under-par dot', 'Daily history chart',
         'Solid green circle. Tiny — used as a chart marker. Could just be '
         'a styled SVG circle with theme color, no PNG needed.'],
        ['🔴', 'Over-par dot', 'Daily history chart',
         'Solid red circle. Same as above.'],
    ])

    # ── Tier 7 ─────────────────────────────────────────
    heading(doc, 'Tier 7 — Achievement-specific icons')
    body(doc, 'src/logic/achievements.js defines ~10. Same 128x128 class as '
              'power-ups. Several already overlap with other tiers (🏆, 🔥, 🛡️, '
              '⛏️, ⏱️) — those reuse Tier 5/6 sprites. New ones below.')
    add_table(doc, [
        ['Icon', 'Achievement', 'Theme', 'Idea direction'],
        ['⚡', 'Lightning / speed', 'Speed-tier achievement',
         'Lightning bolt. Could be diagonal slash style. Distinct from Tier 6 '
         'icons.'],
        ['📅', 'Daily completion',
         'Daily achievement (could overlap with Daily mode card)',
         'A calendar page with a check. Or 7 cells in a row. Could overlap '
         'with the Daily mode-card sprite to save work.'],
        ['📆', 'Weekly completion', 'Weekly streak achievement',
         'A 7-day calendar block. Or a "WK" badge. Differentiate from the daily '
         'calendar sprite.'],
        ['🎮', 'Playtime', 'Hours-played achievement',
         'A controller silhouette. Or an hourglass with cells inside. "Playtime" '
         'could be reframed as "dedication."'],
        ['💪', 'Difficulty', 'High-level achievement',
         'A flexed arm. Or a mountain summit. Or a stack of completed boards. '
         'Should feel "you grinded for this."'],
        ['🎪', 'Chaos', 'Chaos-mode achievement (overlaps with Chaos mode card)',
         'Could reuse the Chaos mode-card sprite. Or a "circus tent" if you '
         'want it distinct — but the visual conflict is OK.'],
    ])

    # ── Tier 8 ─────────────────────────────────────────
    heading(doc, 'Tier 8 — Optional collection theme packs (later)')
    body(doc, 'Each pack overrides mine/flag/smiley sprites with themed alternates. '
              'Lower priority since unlock content. Six packs × 5 sprites = 30 '
              'sprites total. Ship Tier 1 default first; draw these incrementally. '
              'Keep visual coherence WITHIN each pack — all five sprites should '
              'feel like they belong together.')

    pack_table = [
        ['Pack', 'Mine', 'Flag', 'Idle', 'Win', 'Loss'],
        ['Pirate', 'Skull (💀)', 'Pirate flag (🏴‍☠️)', 'Parrot (🦜)',
         'Pirate flag (🏴‍☠️)', 'Skull and crossbones (☠️)'],
        ['Space', 'Comet (☄️)', 'Satellite (🛰️)', 'Astronaut (👨‍🚀)',
         'Rocket (🚀)', 'Explosion (💥)'],
        ['Garden', 'Worm/grub (🐛)', 'Sprout (🌱)', 'Sunflower (🌻)',
         'Cherry blossom (🌸)', 'Wilted flower (🥀)'],
        ['Ocean', 'Shark (🦈)', 'Anchor (⚓)', 'Fish (🐠)',
         'Dolphin (🐬)', 'Bubble (🫧)'],
        ['Medieval', 'Dragon (🐉)', 'Crossed swords (⚔️)', 'Knight\'s shield (🛡️)',
         'Crown (👑)', 'Skull (💀)'],
        ['Holiday', 'Wrapped gift (🎁)', 'Christmas tree (🎄)', 'Santa (🎅)',
         'Star (⭐)', 'Snowman (☃️)'],
    ]
    table = doc.add_table(rows=len(pack_table), cols=6)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(pack_table):
        row = table.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                shade_cell(cell, '2E2E5C')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

    body(doc, 'Pack-level idea direction:')
    body(doc, 'Pirate — woodcut / vintage map aesthetic, sepia palette, '
              'parchment textures.')
    body(doc, 'Space — clean vector, neon palette, starfield backgrounds.')
    body(doc, 'Garden — soft watercolor, pastel palette, organic edges.')
    body(doc, 'Ocean — deep-sea pressure aesthetic, blues and bioluminescent '
              'highlights. Tap into your marine biology angle here — could draw '
              'real species (Atlantic cod, blue shark, lion\'s mane jellyfish).')
    body(doc, 'Medieval — illuminated-manuscript style, gold leaf accents, '
              'gothic line work.')
    body(doc, 'Holiday — warm, textured, snow-flecked. Standard winter palette.')

    # ── Recommended order ─────────────────────────────
    heading(doc, 'Recommended drawing order (effort vs impact)')
    rec_lines = [
        '1. Tier 1 (6 sprites) — biggest visible win for time invested.',
        '2. Tier 5 (4 sprites) — easy as a coherent set, polishes end-of-game flow.',
        '3. Tier 3 (7 sprites) — sets the tone of the title screen.',
        '4. Tier 4 (6 sprites) — power-up bar feels custom now.',
        '5. Tier 2 (10 sprites) — modifier polish.',
        '6. Tier 6 + 7 (~25 sprites) — nav and achievement details.',
    ]
    for line in rec_lines:
        body(doc, line)
    body(doc, '')
    body(doc, 'About 33 sprites for the core polish pass before any theme packs. '
              'At 2-3 sprites per evening, ~2 weeks of drawing.')

    # ── Notes on file naming ──────────────────────────
    heading(doc, 'File naming convention (when ready to drop in)')
    body(doc, 'Save each sprite to /assets/sprites/ with a kebab-case slug:')
    body(doc, '  • Cell sprites: mine.png, flag.png, smiley-idle.png, '
              'smiley-win.png, smiley-loss.png, strike.png')
    body(doc, '  • Modifiers: mod-wall.png, mod-liar.png, mod-mystery.png, '
              'mod-locked.png, mod-wormhole.png, mod-mirror.png, mod-plate.png, '
              'mod-sonar.png, mod-compass.png, mod-mineshift.png')
    body(doc, '  • Modes: mode-challenge.png, mode-quickplay.png, mode-daily.png, '
              'mode-bonus.png, mode-weekly.png, mode-chaos.png, mode-trainer.png')
    body(doc, '  • Power-ups: pow-revealsafe.png, pow-shield.png, pow-scan.png, '
              'pow-lifeline.png, pow-magnet.png, pow-xray.png')
    body(doc, '  • Medals: medal-diamond.png, medal-gold.png, medal-silver.png, '
              'medal-bronze.png')
    body(doc, '  • Nav: nav-home.png, nav-stats.png, etc.')
    body(doc, '  • Theme packs: theme-pirate-mine.png, theme-pirate-flag.png, etc.')
    body(doc, '')
    body(doc, 'When the asset-loading layer is built, we\'ll wire it so missing '
              'sprite files automatically fall back to the current emoji — meaning '
              'you can drop in sprites one at a time without breaking anything '
              'until the full set is done.')

    # Save
    out_path = 'C:/Users/Christopher Wells/OneDrive - Bowdoin College/App Development/Minesweeper/SPRITE_INVENTORY.docx'
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    main()
